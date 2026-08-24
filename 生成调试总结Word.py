#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""生成《归因分析 Agent 调试工作总结.docx》"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ---------- 基础样式 ----------
style = doc.styles["Normal"]
style.font.name = "宋体"
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5


def h1(text):
    p = doc.add_heading(text, level=1)
    return p


def h2(text):
    doc.add_heading(text, level=2)


def para(text, bold=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def numbered(text):
    doc.add_paragraph(text, style="List Number")


# ---------- 封面标题 ----------
t = doc.add_heading("归因分析 Agent 调试工作总结", level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER

para("负责人：徐媛媛（归因分析 Agent）", bold=True)
para("日期：2026-08-24")
para("演示目标：周二下午 · 离线数据 000004.SZ · 全页面正常可用")
para("文档说明：本文件汇总归因分析 Agent 自框架搭建以来的全部调试改动、验证结果与提交记录。", color=(128, 0, 0))
doc.add_paragraph()

# ---------- 一、任务背景 ----------
h1("一、任务背景")
para("团队 7-Agent 扫雷预警系统进入调试阶段，模块分工如下：")
bullet("公告研读 Agent —— 陈吉")
bullet("财务异常 Agent —— 子萍")
bullet("预测建模 Agent —— 万志")
bullet("案例匹配 Agent —— 子鹏")
bullet("归因分析 Agent —— 徐媛媛（本人）")
para("归因分析 Agent 的职责：把预测模型的概率与 SHAP 特征贡献，翻译成「可读风险因素 + 原文证据 + 相似案例」三元组，回答『为什么预测这家公司会被问询』。", bold=True)

# ---------- 二、调试发现的问题 ----------
h1("二、调试发现的问题（共 5 项）")

h2("问题 1：特征映射表覆盖严重不足（最关键）")
para("归因 Agent 的 FEATURE_MAP 只覆盖 14 个模型特征，而预测模型 models_manifest.json 共 135 个特征。")
bullet("89%（121 个）特征映射后显示为「其他/未知」，SHAP 归因基本失效")
bullet("涉及 f2/f6/gov/mkt/sent/regulatory_inquiry_semantic 六大类")
bullet("50 个问询函语义嵌入维度（regulatory_inquiry_semantic_*）无法逐条命名")

h2("问题 2：无 LLM 叙事时证据引用为空")
para("evidence_citations 只保留「叙事中引用的证据」，use_llm=False 时叙事为空 → 页面「证据池」空白。")

h2("问题 3：找不到证据的 SHAP 因素被直接剔除")
para("evidence_locate 只保留有证据绑定的因素，实测 8 个 SHAP 特征仅 1 个存活，演示观感极差。")

h2("问题 4：归因页漏了案例匹配环节")
para("analyze_company 跑了「公告研读→财务→预测→归因」，但漏跑 CaseRetrieverAgent，导致归因的「相似案例链接」永远为空。")

h2("问题 5：环境与数据路径")
bullet("streamlit / langgraph / PyMuPDF / pdfplumber / python-dotenv 未安装")
bullet("公告 PDF 实际在 C 盘（上市公司公告与定期报告数据集），配置默认指向不存在的 D:\\BaiduNetdiskDownload")
bullet("公告索引缓存（000004_SZ_index.json）仍是旧 D 盘路径")

# ---------- 三、代码修改明细 ----------
h1("三、代码修改明细")

h2("1. backend/agents/attributor.py（归因核心逻辑）")
numbered("FEATURE_MAP 从 14 个扩到 135 个：补齐 f2(12)/f6(11)/gov(13)/mkt(28)/sent(19)/基础指标，全部给出中文描述与标签引用")
numbered("新增 PREFIX_MAP 前缀降级：50 个语义嵌入维度归为「语义信号」组，未命中特征按前缀归类，消除「其他/未知」")
numbered("新增 _resolve_feature()：精确匹配 → 前缀匹配 → 未知 三级解析")
numbered("map_factors() 统一 factor 结构（description / evidence_id / is_fallback）")
numbered("evidence_locate()：无证据因素不再剔除，标记 no_evidence=True（模型侧信号）")
numbered("evidence_citations 改为「诱因绑定证据 + 叙事引用证据」的白名单过滤，无 LLM 时不再为空")
numbered("新增 _factor_tag() 展示标签：SHAP=xx/规则降级 + 证据状态")

h2("2. 归因分析agent.py（Streamlit 页面）")
numbered("analyze_company 尊重 ANNOUNCE_SOURCE=local 离线模式（离线时不再硬编码 Cninfo 在线源）")
numbered("补上 CaseRetrieverAgent 环节（案例链接不再为空）")
numbered("新增「相似历史案例」展示区 + 指标卡")
numbered("factors_dataframe 对无证据因素显示「无直接证据」")

h2("3. backend/tests/test_attributor.py（单元测试）")
numbered("更新 evidence_citations 断言（无 LLM 时含诱因绑定证据）")

h2("4. .env（本地配置，不入库）")
bullet("DATA_RAW → C:/Users/18316/Desktop/比赛/金融比赛/上市公司公告与定期报告数据集")
bullet("ANNOUNCE_SOURCE=local（离线 PDF 扫描）")
bullet("ANNUAL_REPORT_DIR / INQUIRY_DATA_DIR 指向实际数据位置")
bullet("EMBEDDING_BACKEND=fallback（未装 torch 时兜底）")

# ---------- 四、测试与验证 ----------
h1("四、测试与验证结果")
bullet("单元测试：31 个用例全部通过（backend/tests/）")
bullet("完整离线流水线（7 节点 LangGraph）全部 done：公告研读→财务→预测→案例→chunk→归因→报告")
bullet("000004.SZ 实测：13 份公告、28 风险要素、财务 3 异常（高）、60d 概率 0.117、归因 8 诱因、3 相似案例、3 证据引用，总耗时约 18 秒")
bullet("Streamlit 页面正常启动（http://localhost:8506，HTTP 200）")
bullet("已安装依赖：streamlit 1.62 / langgraph / PyMuPDF 1.28 / pdfplumber / python-dotenv / lightgbm / xgboost / shap")

# ---------- 五、Git 提交记录 ----------
h1("五、Git 提交记录（分支 feature/attribution-agent-xyy，已推送 origin）")
bullet("94e645f  fix(attributor): 归因页补上案例匹配环节 + 展示相似案例")
bullet("82f257e  fix(attributor): 归因页尊重 ANNOUNCE_SOURCE=local 离线模式")
bullet("0147d7b  fix(attributor): 补全特征映射表 + 优化证据定位")
para("改动文件：backend/agents/attributor.py、归因分析agent.py、backend/tests/test_attributor.py（均为归因模块，不影响其他 Agent）", bold=True)

# ---------- 六、演示步骤 ----------
h1("六、演示步骤（离线 000004.SZ）")
numbered("cd E:\\agent")
numbered("python -m streamlit run 归因分析agent.py --server.port 8506")
numbered("浏览器打开 http://localhost:8506")
numbered("输入 000004.SZ → 点「开始归因」→ 等待约 20-60 秒")
numbered("展示：概率/风险等级 → Top 8 诱因（SHAP+证据）→ 相似案例 → 审计追踪")

# ---------- 七、遗留事项 ----------
h1("七、遗留事项与建议")
bullet("torch 未安装：案例匹配 BGE 语义向量通道未启用（当前走标签通道），如需启用由案例匹配负责人协调安装（约 2GB）")
bullet(".env 与公告索引为本地机器相关配置，未提交（避免影响队友的 D 盘路径）")
bullet("演示话术建议：强调「特征映射 + 证据白名单 + 防幻觉校验」三个亮点")

doc.save("归因分析Agent调试总结.docx")
print("已生成：归因分析Agent调试总结.docx")
